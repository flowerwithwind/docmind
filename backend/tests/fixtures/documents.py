"""测试文档构建器：生成内存中的 docx / xlsx / pdf 字节流。"""
from __future__ import annotations

import io

import docx
from openpyxl import Workbook
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas


def make_docx_bytes(paragraphs: list[tuple[str | None, str]]) -> bytes:
    """paragraphs: [(heading_level | None, text)]。"""
    doc = docx.Document()
    for level, text in paragraphs:
        if level:
            doc.add_heading(text, level=level)
        else:
            doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def make_docx_with_table_bytes() -> bytes:
    doc = docx.Document()
    doc.add_heading("购销合同", level=1)
    doc.add_paragraph("甲乙双方就以下条款达成一致。")
    table = doc.add_table(rows=3, cols=2)
    cells = [("品名", "数量"), ("服务器", "10 台"), ("交换机", "5 台")]
    for i, (a, b) in enumerate(cells):
        table.rows[i].cells[0].text = a
        table.rows[i].cells[1].text = b
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def make_xlsx_bytes(rows: list[list[str]]) -> bytes:
    wb = Workbook()
    ws = wb.active
    ws.title = "Sheet1"
    for r in rows:
        ws.append(r)
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def make_mini_pdf_bytes(lines: list[str]) -> bytes:
    """用 reportlab 生成一页 PDF。"""
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    c.setFont("Helvetica", 14)
    y = 780
    for i, line in enumerate(lines):
        c.drawString(72, y - i * 20, line)
    c.save()
    return buf.getvalue()
