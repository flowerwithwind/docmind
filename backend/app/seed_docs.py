"""演示样例文档构建（M5/FR-09）：内存生成 docx，零外部资产依赖。

与 tests/helpers.py 的段落保持一致，确保规则抽取器可命中；
contract_v2 为"合同对比版"，金额/交付期限/违约金与 v1 不同，用于 M5 对比演示。
"""
from __future__ import annotations

import io

import docx

CONTRACT_PARAGRAPHS: list[tuple[int | None, str]] = [
    (1, "产品购销合同"),
    (2, "第一条 当事人"),
    (None, "合同名称：产品购销合同"),
    (None, "甲方：华信科技有限公司"),
    (None, "乙方：远景供应链有限公司"),
    (None, "签订日期：2026年3月15日"),
    (None, "合同金额：人民币壹佰贰拾万元整，币种为人民币。"),
    (2, "第二条 付款与交付"),
    (None, "付款方式：预付款30%，验收合格后70%。"),
    (None, "交付期限：合同签订后30日内交货。"),
    (None, "违约金条款：逾期每日按合同金额的0.5%支付违约金。"),
    (None, "争议解决：双方争议提交甲方所在地人民法院诉讼解决。"),
    (None, "有效期：本合同自签订之日起一年。"),
]

CONTRACT_V2_PARAGRAPHS: list[tuple[int | None, str]] = [
    (1, "产品购销合同"),
    (2, "第一条 当事人"),
    (None, "合同名称：产品购销合同"),
    (None, "甲方：华信科技有限公司"),
    (None, "乙方：远景供应链有限公司"),
    (None, "签订日期：2026年6月1日"),
    (None, "合同金额：人民币壹佰叁拾万元整，币种为人民币。"),
    (2, "第二条 付款与交付"),
    (None, "付款方式：预付款30%，验收合格后70%。"),
    (None, "交付期限：合同签订后45日内交货。"),
    (None, "违约金条款：逾期每日按合同金额的0.8%支付违约金。"),
    (None, "争议解决：双方争议提交甲方所在地人民法院诉讼解决。"),
    (None, "有效期：本合同自签订之日起两年。"),
]

FINANCIAL_PARAGRAPHS: list[tuple[int | None, str]] = [
    (1, "2025年年度报告"),
    (None, "报告期：2025年年度，公司主营业务：智能硬件研发与销售。"),
    (None, "营业收入85600万元，营业收入同比增长12.5%。"),
    (None, "净利润9800万元，净利润同比增长8.2%。"),
    (None, "毛利率32.1%，资产负债率45.6%，基本每股收益1.25元。"),
]

BUILDERS = {
    "contract": ("demo-contract.docx", CONTRACT_PARAGRAPHS),
    "contract_v2": ("demo-contract-v2.docx", CONTRACT_V2_PARAGRAPHS),
    "financial": ("demo-financial.docx", FINANCIAL_PARAGRAPHS),
}


def make_demo_docx(kind: str) -> bytes:
    """按 kind 生成 docx 字节流；kind 非法时抛 KeyError。"""
    _, paragraphs = BUILDERS[kind]
    doc = docx.Document()
    for level, text in paragraphs:
        if level:
            doc.add_heading(text, level=level)
        else:
            doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def demo_filename(kind: str) -> str:
    return BUILDERS[kind][0]
